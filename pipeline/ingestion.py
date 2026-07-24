import os
from pathlib import Path
from typing import List, Optional
from models import ResumeDocument
from utils.logger import logger

def extract_text_from_txt(filepath: str) -> str:
    """Extracts text from a plain text file using UTF-8 / latin-1 fallback."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        with open(filepath, "r", encoding="latin-1") as f:
            return f.read().strip()

def extract_text_from_pdf(filepath: str) -> str:
    """Extracts text from a PDF file using pdfplumber with pypdf fallback."""
    text_chunks = []
    
    # Try pdfplumber with layout-aware column preservation first
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                # Try layout-aware extraction first for multi-column resumes
                page_text = page.extract_text(layout=True)
                if not page_text or len(page_text.strip()) < 10:
                    page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
        extracted = "\n".join(text_chunks).strip()
        if extracted:
            return extracted
    except Exception as e:
        logger.debug(f"pdfplumber failed for {filepath}, trying pypdf fallback: {e}")

    # Fallback to pypdf / PyPDF2
    try:
        try:
            import pypdf as pdf_lib
        except ImportError:
            import PyPDF2 as pdf_lib  # fallback import name
            
        reader = pdf_lib.PdfReader(filepath)
        if reader.is_encrypted:
            # Try empty password
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError("PDF is password protected / encrypted")
                
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
        return "\n".join(text_chunks).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(filepath: str) -> str:
    """Extracts text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        return "\n".join(full_text).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

def load_single_file(filepath: str) -> ResumeDocument:
    """Safely loads and extracts text from a single file with per-file try/except error handling."""
    path = Path(filepath)
    ext = path.suffix.lower()
    filename = path.name

    doc = ResumeDocument(
        filename=filename,
        filepath=str(path.absolute()),
        extension=ext,
        raw_text="",
        extraction_status="success",
        extraction_error=None
    )

    if not path.exists():
        doc.extraction_status = "unreadable_file"
        doc.extraction_error = "File does not exist"
        return doc

    try:
        if ext == ".txt":
            raw_text = extract_text_from_txt(filepath)
        elif ext == ".pdf":
            raw_text = extract_text_from_pdf(filepath)
        elif ext in [".docx", ".doc"]:
            raw_text = extract_text_from_docx(filepath)
        else:
            # Unsupported extension or binary
            raw_text = extract_text_from_txt(filepath)

        if not raw_text or len(raw_text.strip()) < 10:
            logger.warning(f"File {filename} produced minimal or empty text.")
            # Scanned/image-only PDF without OCR yields empty text
            if ext == ".pdf":
                doc.extraction_status = "unreadable_file"
                doc.extraction_error = "PDF yields no extractable text (likely image-only/scanned or empty)"
                return doc

        doc.raw_text = raw_text
        return doc

    except Exception as e:
        logger.error(f"Error loading file {filename}: {e}")
        doc.extraction_status = "unreadable_file"
        doc.extraction_error = str(e)
        return doc

def load_job_description(jd_path: str) -> str:
    """Loads text from JD file (supports TXT, PDF, DOCX)."""
    logger.info(f"Loading Job Description from: {jd_path}")
    doc = load_single_file(jd_path)
    if doc.extraction_status != "success" or not doc.raw_text:
        raise ValueError(f"Could not extract text from Job Description at {jd_path}: {doc.extraction_error}")
    return doc.raw_text

def load_resumes_from_directory(folder_path: str) -> List[ResumeDocument]:
    """Loads all resume files from a target directory with per-file exception handling."""
    logger.info(f"Scanning resume directory: {folder_path}")
    folder = Path(folder_path)
    if not folder.exists() or not folder.is_dir():
        raise ValueError(f"Resume directory does not exist or is not a directory: {folder_path}")

    documents = []
    # Sort files for deterministic batch execution
    files = sorted([f for f in folder.iterdir() if f.is_file() and not f.name.startswith(".")])
    
    for file_path in files:
        doc = load_single_file(str(file_path))
        documents.append(doc)
        logger.info(f"Loaded {doc.filename} | Status: {doc.extraction_status} | Length: {len(doc.raw_text)} chars")

    return documents
